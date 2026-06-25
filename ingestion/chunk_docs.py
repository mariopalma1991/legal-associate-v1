"""
Extracts text from downloaded documents (PDF, DOCX, XLSX) and splits into chunks.

Reads:   documents WHERE status = 'downloaded' (Vigente licitaciones only)
Writes:  chunks table (id, licitacion_id, document_id, doc_type,
                       chunk_index, page_number, text, token_count)
Updates: documents SET status = 'chunked'

Usage:
  python chunk_docs.py              # chunk all downloaded docs
  python chunk_docs.py --limit 10  # chunk N docs (testing)
  python chunk_docs.py --chunk-size 256 --overlap 30
"""

import argparse
import json
import os
import time
import uuid

import fitz          # PyMuPDF
import tiktoken
import psycopg2
import psycopg2.extras
from dotenv import load_dotenv
import sys as _sys; _sys.path.insert(0, __import__('os').path.dirname(__import__('os').path.dirname(__import__('os').path.abspath(__file__))))
from shared._storage import download_doc, local_to_key
from _pipeline import start_run, finish_run, start_stage, finish_stage

load_dotenv()

CHUNK_SIZE = 1024   # tokens per chunk
OVERLAP    = 256    # token overlap between consecutive chunks
BATCH_SIZE = 50


# â”€â”€ DB connection â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def get_conn():
    url = os.getenv("DATABASE_URL")
    if not url:
        raise ValueError("DATABASE_URL not set in .env")
    at = url.rfind("@")
    creds, host_part = url[len("postgresql://"):at], url[at + 1:]
    user, password = creds.split(":", 1)
    host, port = host_part.split("/")[0].split(":")
    dbname = host_part.split("/")[1].split("?")[0]
    return psycopg2.connect(
        host=host, port=int(port), dbname=dbname,
        user=user, password=password, sslmode="require"
    )


# â”€â”€ Raw-text cache â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def _load_raw_text(conn, doc_id: str) -> list[tuple[int, str]] | None:
    """Return cached [(page_num, text)] from documents.raw_text, or None."""
    cur = conn.cursor()
    cur.execute("SELECT raw_text FROM documents WHERE id = %s", (doc_id,))
    row = cur.fetchone()
    cur.close()
    if row and row[0]:
        return [(entry["page"], entry["text"]) for entry in row[0]]
    return None


def _save_raw_text(conn, doc_id: str, pages: list[tuple[int, str]]):
    """Write [(page_num, text)] to documents.raw_text (runs inside caller's txn)."""
    data = [{"page": p, "text": t} for p, t in pages]
    cur = conn.cursor()
    cur.execute(
        "UPDATE documents SET raw_text = %s WHERE id = %s",
        (json.dumps(data), doc_id),
    )
    cur.close()


# â”€â”€ Chunking â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def split_tokens(text: str, enc, chunk_size: int, overlap: int) -> list[str]:
    tokens = enc.encode(text)
    if not tokens:
        return []
    result = []
    start = 0
    while start < len(tokens):
        end = min(start + chunk_size, len(tokens))
        result.append(enc.decode(tokens[start:end]))
        if end == len(tokens):
            break
        start += chunk_size - overlap
    return result


def _build_chunks(texts_with_pages: list[tuple[int, str]], doc_id, licitacion_id,
                  tipo, enc, chunk_size, overlap) -> list[dict]:
    """Turn (page_num, text) pairs into chunk dicts."""
    chunks = []
    chunk_index = 0
    for page_num, text in texts_with_pages:
        for chunk_text in split_tokens(text, enc, chunk_size, overlap):
            chunk_text = chunk_text.strip()
            if not chunk_text:
                continue
            chunks.append({
                "id":            str(uuid.uuid4()),
                "licitacion_id": licitacion_id,
                "document_id":   doc_id,
                "doc_type":      tipo,
                "chunk_index":   chunk_index,
                "page_number":   page_num,
                "text":          chunk_text,
                "token_count":   len(enc.encode(chunk_text)),
            })
            chunk_index += 1
    return chunks


def _extract_pdf_pymupdf(local_path: str) -> tuple[list[tuple[int, str]], str | None]:
    try:
        pdf = fitz.open(local_path)
        pages = [(i + 1, pdf[i].get_text().strip()) for i in range(pdf.page_count)]
        pdf.close()
        return [(p, t) for p, t in pages if t], None
    except Exception as e:
        return [], str(e)


def _extract_pdf_llamaparse(local_path: str) -> tuple[list[tuple[int, str]], str | None]:
    try:
        from llama_parse import LlamaParse
        parser = LlamaParse(
            api_key=os.getenv("LLAMA_CLOUD_API_KEY"),
            result_type="markdown",
            language="es",
            verbose=False,
        )
        docs = parser.load_data(local_path)
        if not docs:
            return [], "LlamaParse returned no content"
        # Each doc is one page; join all into sequential pages
        pages = [(i + 1, doc.text.strip()) for i, doc in enumerate(docs) if doc.text.strip()]
        if not pages:
            return [], "LlamaParse returned empty text"
        return pages, None
    except Exception as e:
        return [], str(e)


def _extract_pdf(local_path: str, use_llamaparse: bool = True) -> tuple[list[tuple[int, str]], str | None]:
    if use_llamaparse:
        pages, error = _extract_pdf_llamaparse(local_path)
        if pages:
            return pages, None
        # Fallback to PyMuPDF if LlamaParse fails or returns nothing
        print(f"    [LlamaParse fallback] {error} â€” trying PyMuPDF")
        return _extract_pdf_pymupdf(local_path)
    return _extract_pdf_pymupdf(local_path)


def _extract_zip(local_path: str, use_llamaparse: bool = True) -> tuple[list[tuple[int, str]], str | None]:
    """Extract all PDFs, DOCX, and XLSX files inside a ZIP."""
    import zipfile, tempfile
    SUPPORTED = (".pdf", ".docx", ".xlsx")
    try:
        pages = []
        global_page = 1
        with zipfile.ZipFile(local_path, "r") as zf:
            entries = [
                n for n in zf.namelist()
                if os.path.splitext(n.lower())[1] in SUPPORTED
                and not os.path.basename(n).startswith("__")
            ]
            if not entries:
                return [], "ZIP contains no supported files (pdf/docx/xlsx)"
            for entry in sorted(entries):
                ext  = os.path.splitext(entry.lower())[1]
                data = zf.read(entry)
                with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                    tmp.write(data)
                    tmp_path = tmp.name
                try:
                    if ext == ".pdf":
                        entry_pages, _ = _extract_pdf(tmp_path, use_llamaparse=use_llamaparse)
                    elif ext == ".docx":
                        entry_pages, _ = _extract_docx(tmp_path)
                    elif ext == ".xlsx":
                        entry_pages, _ = _extract_xlsx(tmp_path)
                    else:
                        entry_pages = []
                finally:
                    os.unlink(tmp_path)
                for _, text in entry_pages:
                    pages.append((global_page, text))
                    global_page += 1
        return pages, None
    except Exception as e:
        return [], str(e)


def _extract_rar(local_path: str, use_llamaparse: bool = True) -> tuple[list[tuple[int, str]], str | None]:
    """Extract all PDFs, DOCX, and XLSX files inside a RAR archive."""
    import rarfile, tempfile
    rarfile.UNRAR_TOOL = r"C:\Program Files\WinRAR\UnRAR.exe"
    SUPPORTED = (".pdf", ".docx", ".xlsx")
    try:
        pages = []
        global_page = 1
        with rarfile.RarFile(local_path, "r") as rf:
            entries = [
                n for n in rf.namelist()
                if os.path.splitext(n.lower())[1] in SUPPORTED
                and not os.path.basename(n).startswith("__")
            ]
            if not entries:
                return [], "RAR contains no supported files (pdf/docx/xlsx)"
            for entry in sorted(entries):
                ext  = os.path.splitext(entry.lower())[1]
                data = rf.read(entry)
                with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                    tmp.write(data)
                    tmp_path = tmp.name
                try:
                    if ext == ".pdf":
                        entry_pages, _ = _extract_pdf(tmp_path, use_llamaparse=use_llamaparse)
                    elif ext == ".docx":
                        entry_pages, _ = _extract_docx(tmp_path)
                    elif ext == ".xlsx":
                        entry_pages, _ = _extract_xlsx(tmp_path)
                    else:
                        entry_pages = []
                finally:
                    os.unlink(tmp_path)
                for _, text in entry_pages:
                    pages.append((global_page, text))
                    global_page += 1
        return pages, None
    except Exception as e:
        return [], str(e)


def _extract_docx(local_path: str) -> tuple[list[tuple[int, str]], str | None]:
    try:
        from docx import Document
        doc = Document(local_path)
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    lines.append(row_text)
        return [(1, "\n".join(lines))] if lines else [], None
    except Exception as e:
        return [], str(e)


def _extract_xlsx(local_path: str) -> tuple[list[tuple[int, str]], str | None]:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(local_path, read_only=True, data_only=True)
        pages = []
        for sheet_num, ws in enumerate(wb.worksheets, 1):
            rows = []
            for row in ws.iter_rows(values_only=True):
                vals = [str(v) for v in row if v is not None and str(v).strip()]
                if vals:
                    rows.append(" | ".join(vals))
            if rows:
                pages.append((sheet_num, f"[Hoja: {ws.title}]\n" + "\n".join(rows)))
        wb.close()
        return pages, None
    except Exception as e:
        return [], str(e)


def extract_chunks(doc_id: str, licitacion_id: int, tipo: str,
                   local_path: str, enc,
                   chunk_size: int, overlap: int,
                   use_llamaparse: bool = True,
                   conn=None) -> tuple[list[dict], str | None]:
    """
    Dispatch to the right extractor based on file extension.

    If conn is provided, checks documents.raw_text first â€” if cached, skips
    extraction entirely.  After a fresh extraction the result is written to
    raw_text (committed atomically with write_chunks' transaction).
    """
    ext = os.path.splitext(local_path)[1].lower()

    # â”€â”€ Cache hit â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if conn:
        pages = _load_raw_text(conn, doc_id)
        if pages is not None:
            return _build_chunks(pages, doc_id, licitacion_id, tipo, enc, chunk_size, overlap), None

    # â”€â”€ Fresh extraction â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    if ext == ".pdf":
        pages, error = _extract_pdf(local_path, use_llamaparse=use_llamaparse)
    elif ext == ".docx":
        pages, error = _extract_docx(local_path)
    elif ext == ".xlsx":
        pages, error = _extract_xlsx(local_path)
    elif ext == ".zip":
        pages, error = _extract_zip(local_path, use_llamaparse=use_llamaparse)
    elif ext == ".rar":
        pages, error = _extract_rar(local_path, use_llamaparse=use_llamaparse)
    else:
        return [], f"Unsupported file type: {ext}"

    if error:
        return [], error

    # Cache for future re-chunking (committed inside write_chunks' transaction)
    if conn and pages:
        _save_raw_text(conn, doc_id, pages)

    return _build_chunks(pages, doc_id, licitacion_id, tipo, enc, chunk_size, overlap), None


# â”€â”€ DB writes â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def write_chunks(conn, doc_id: str, chunks: list[dict], error: str | None = None,
                 parser_name: str = "pymupdf", chunk_config: str = "1024_256"):
    """
    In a single transaction:
      - Delete existing chunks for this document+config only (other configs preserved)
      - Insert new chunks tagged with chunk_config
      - Mark document as 'chunked' (or 'error')
    """
    with conn:
        cur = conn.cursor()
        cur.execute(
            "DELETE FROM chunks WHERE document_id = %s AND chunk_config = %s",
            (doc_id, chunk_config),
        )

        if chunks:
            psycopg2.extras.execute_values(cur, """
                INSERT INTO chunks
                    (id, licitacion_id, document_id, doc_type,
                     chunk_index, page_number, text, token_count, chunk_config)
                VALUES %s
            """, [(c["id"], c["licitacion_id"], c["document_id"], c["doc_type"],
                   c["chunk_index"], c["page_number"], c["text"], c["token_count"],
                   chunk_config)
                  for c in chunks])

        if error:
            cur.execute("""
                UPDATE documents SET status = 'error', error = %s, updated_at = now()
                WHERE id = %s
            """, (error[:200], doc_id))
        else:
            cur.execute("""
                UPDATE documents SET status = 'chunked', parser = %s, updated_at = now()
                WHERE id = %s
            """, (parser_name, doc_id))

        cur.close()


# â”€â”€ Main â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def _process_batch(conn, docs, enc, chunk_size, overlap, use_llamaparse,
                   rechunk_mode: bool, chunk_config: str) -> tuple[int, int, int]:
    """Process a list of (doc_id, licitacion_id, tipo, local_path, parser) rows.
    Returns (processed, errors, total_chunks)."""
    processed = errors = total_chunks = 0
    total = len(docs)
    t0 = time.time()

    for i, row in enumerate(docs, 1):
        doc_id, licitacion_id, tipo, local_path, stored_parser = row

        tmp_path = None
        if local_path and os.path.exists(local_path):
            effective_path = local_path
        else:
            # File is in Supabase Storage â€” download to a tempfile
            import tempfile
            key = local_to_key(local_path) if local_path and local_path.startswith("docs") else local_path
            data = download_doc(key) if key else None
            if not data:
                print(f"  [MISS]  id={doc_id}  not found locally or in Storage")
                if not rechunk_mode:
                    write_chunks(conn, str(doc_id), [], error="file not found in Storage",
                                 chunk_config=chunk_config)
                errors += 1
                continue
            ext = os.path.splitext(key)[1] or ".pdf"
            tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
            tmp.write(data)
            tmp.close()
            tmp_path = tmp.name
            effective_path = tmp_path
            print(f"  [FETCH] id={doc_id}  fetched from Storage")

        try:
            chunks, error = extract_chunks(
                str(doc_id), licitacion_id, tipo, effective_path,
                enc, chunk_size, overlap,
                use_llamaparse=use_llamaparse,
                conn=conn,
            )
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

        if error:
            print(f"  [FAIL]  id={doc_id}  {error}")
            if not rechunk_mode:
                write_chunks(conn, str(doc_id), [], error=error,
                             chunk_config=chunk_config)
            errors += 1
            continue

        parser_name = stored_parser if rechunk_mode else ("llamaparse" if use_llamaparse else "pymupdf")
        write_chunks(conn, str(doc_id), chunks,
                     parser_name=parser_name, chunk_config=chunk_config)
        total_chunks += len(chunks)
        processed += 1

        elapsed = time.time() - t0
        rate = processed / elapsed if elapsed > 0 else 0
        eta  = (total - i) / rate / 60 if rate > 0 else 0
        print(f"  [OK]    lid={licitacion_id:<8}  {tipo[:35]:<35}  "
              f"{len(chunks):>3} chunks  "
              f"[{chunk_config}]  "
              f"({i}/{total}  ~{eta:.1f} min remaining)")

    return processed, errors, total_chunks


def main():
    ap = argparse.ArgumentParser(description="Chunk documents into text segments for embedding")
    ap.add_argument("--limit",      type=int, default=None,
                    help="Process only N documents (testing)")
    ap.add_argument("--chunk-size", type=int, default=CHUNK_SIZE,
                    help=f"Tokens per chunk (default: {CHUNK_SIZE})")
    ap.add_argument("--overlap",    type=int, default=OVERLAP,
                    help=f"Overlap tokens between chunks (default: {OVERLAP})")
    ap.add_argument("--parser",     choices=["llamaparse", "pymupdf"], default="llamaparse",
                    help="PDF extractor: llamaparse (default, vision+LLM) or pymupdf (fast, local)")
    ap.add_argument("--rechunk",    action="store_true",
                    help="Re-chunk all docs with cached raw_text using new chunk params "
                         "(no LlamaParse calls â€” instant re-chunking for experimentation)")
    args = ap.parse_args()
    use_llamaparse = args.parser == "llamaparse"

    chunk_config = f"{args.chunk_size}_{args.overlap}"
    enc      = tiktoken.get_encoding("cl100k_base")
    conn     = get_conn()
    run_id   = start_run(conn, notes="chunk")
    stage_id = start_stage(conn, run_id, "chunk", config={
        "chunk_config": chunk_config, "parser": args.parser, "rechunk": args.rechunk,
    })
    cur  = conn.cursor()

    if args.rechunk:
        # â”€â”€ Rechunk mode: collect all docs with cached text â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        print(f"[rechunk] config={chunk_config}  Using cached raw_text")
        cur.execute("""
            SELECT d.id, d.licitacion_id, d.tipo, d.local_path, d.parser
            FROM documents d
            JOIN licitaciones l ON d.licitacion_id = l.id
            WHERE d.raw_text IS NOT NULL
              AND l.licitacion_status = 'Vigente'
            ORDER BY d.licitacion_id, d.tipo
        """)
        all_docs = cur.fetchall()
        cur.close()

        if not all_docs:
            print("No documents with cached raw_text found. Run without --rechunk first.")
            conn.close()
            return

        if args.limit:
            all_docs = all_docs[:args.limit]

        print(f"Found {len(all_docs)} docs with raw_text â€” re-chunking ...")
        t0 = time.time()
        processed, errors, total_chunks = _process_batch(
            conn, all_docs, enc, args.chunk_size, args.overlap,
            use_llamaparse=False, rechunk_mode=True, chunk_config=chunk_config,
        )

    else:
        # â”€â”€ Normal mode: chunk newly downloaded docs â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        print(f"Parser: {args.parser}  config={chunk_config}")
        cur.execute("""
            SELECT COUNT(*) FROM documents d
            JOIN licitaciones l ON d.licitacion_id = l.id
            WHERE d.status = 'downloaded'
              AND l.licitacion_status = 'Vigente'
        """)
        row = cur.fetchone()
        total_available = row[0] if row else 0
        cur.close()

        if total_available == 0:
            print("No downloaded documents to chunk.")
            conn.close()
            return

        total = min(total_available, args.limit) if args.limit else total_available
        print(f"Found {total_available:,} downloaded docs â€” chunking {total:,} ...")

        cur = conn.cursor()
        cur.execute("""
            SELECT d.id, d.licitacion_id, d.tipo, d.local_path, NULL as parser
            FROM documents d
            JOIN licitaciones l ON d.licitacion_id = l.id
            WHERE d.status = 'downloaded'
              AND l.licitacion_status = 'Vigente'
            ORDER BY d.licitacion_id, d.tipo
            LIMIT %s
        """, (total,))
        all_docs = cur.fetchall()
        cur.close()

        t0 = time.time()
        processed, errors, total_chunks = _process_batch(
            conn, all_docs, enc, args.chunk_size, args.overlap,
            use_llamaparse=use_llamaparse, rechunk_mode=False, chunk_config=chunk_config,
        )

    elapsed = time.time() - t0
    print(f"\nDone in {elapsed/60:.1f} min")
    print(f"  Processed : {processed:,} documents")
    print(f"  Chunks    : {total_chunks:,}")
    print(f"  Errors    : {errors}")

    stage_status = "completed" if errors == 0 else ("partial" if processed > 0 else "failed")
    finish_stage(conn, stage_id, stage_status,
                 items_found=processed + errors, items_ok=processed, items_error=errors)
    finish_run(conn, run_id, stage_status)
    conn.close()


if __name__ == "__main__":
    main()
